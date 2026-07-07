#!/usr/bin/env python3
"""
◊ USING THIS MATTER ◊ — hardcover interior builder.

Reads the manuscript (a manifest + chapter files in the estate's micro-markup)
and assembles a 6x9in hardcover-interior DOCX in the estate's typographic canon:
Cormorant Garamond for display, Crimson Pro for prose, IBM Plex Mono for the
procedural voice. The glass is shown: this builder ships with the book.

Micro-markup, one block per stretch of text, blocks separated by blank lines:

  #K text        -> chapter kicker (mono, letterspaced)
  #T text        -> chapter title (display serif)
  #E text        -> epigraph line (italic, centered-ish block)
  #ES text       -> epigraph source (small caps line under epigraph)
  ::: doc        -> begin mono document block (until ':::')
  ::: letter     -> begin letter block, serif italic (until ':::')
  ::: annotation -> begin second-hand marginalia block (until ':::')
  ::: verse      -> begin centered verse block (until ':::')
  @sep           -> ◊ ◈ ◊ separator
  @break         -> ·  ·  · soft break
  @blank         -> vertical blank
  @fig:file.png|caption        -> figure with caption
  @plate:file.png|caption      -> full-page plate (page break before + after)
  *i* and **b**  -> inline italic / bold
  plain text     -> body paragraph (first after heading/sep is unindented)
"""

import json
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
BOOK = os.path.dirname(HERE)

SERIF_BODY = "Crimson Pro"
SERIF_DISPLAY = "Cormorant Garamond"
MONO = "IBM Plex Mono"

INK = RGBColor(0x1A, 0x18, 0x15)        # near-black warm ink
ASH = RGBColor(0x6B, 0x6B, 0x6B)        # the witness's register
EMBER = RGBColor(0x8A, 0x6D, 0x1B)      # ember gold, darkened for paper
BLOOD = RGBColor(0x7A, 0x1F, 0x1F)      # used sparingly

BODY_PT = 11
BODY_LEADING = 15.6


# ---------------------------------------------------------------- utilities

def set_font(run, name=SERIF_BODY, size=BODY_PT, color=INK, bold=False,
             italic=False, caps=False, small_caps=False, spacing=None):
    f = run.font
    f.name = name
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), name)
    r.set(qn("w:cs"), name)
    f.size = Pt(size)
    f.color.rgb = color
    f.bold = bold
    f.italic = italic
    if caps:
        f.all_caps = True
    if small_caps:
        f.small_caps = True
    if spacing is not None:  # letterspacing in twentieths of a point
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(spacing * 20)))
        run._element.rPr.append(sp)
    return run


def para(doc, align=None, before=0, after=0, leading=None, indent=None,
         left=None, right=None, keep_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if leading is not None:
        pf.line_spacing = Pt(leading)
    if indent is not None:
        pf.first_line_indent = Inches(indent)
    if left is not None:
        pf.left_indent = Inches(left)
    if right is not None:
        pf.right_indent = Inches(right)
    if keep_next:
        pf.keep_with_next = True
    return p


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline(p, text, name=SERIF_BODY, size=BODY_PT, color=INK,
               base_italic=False, base_bold=False, **kw):
    """Add text with *italic* and **bold** inline markup."""
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            set_font(p.add_run(tok[2:-2]), name=name, size=size, color=color,
                     bold=True, italic=base_italic, **kw)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            set_font(p.add_run(tok[1:-1]), name=name, size=size, color=color,
                     italic=not base_italic, bold=base_bold, **kw)
        else:
            set_font(p.add_run(tok), name=name, size=size, color=color,
                     italic=base_italic, bold=base_bold, **kw)
    return p


def page_field(p, name=MONO, size=8.5, color=ASH):
    r1 = p.add_run()
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    r1._element.append(fld)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r2._element.append(instr)
    r3 = p.add_run()
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    r3._element.append(fld2)
    for r in (r1, r2, r3):
        set_font(r, name=name, size=size, color=color)


def clear_para(container):
    """Return the first (empty) paragraph of a header/footer, cleared."""
    p = container.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    return p


def sect_pgnum(section, fmt=None, start=None):
    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(el)
    el = OxmlElement("w:pgNumType")
    if fmt:
        el.set(qn("w:fmt"), fmt)
    if start is not None:
        el.set(qn("w:start"), str(start))
    sectPr.append(el)


def setup_section(section, first_page_blank_header=True):
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.42)


def running_heads(section, verso_text, recto_text, show_numbers=True):
    """Even = verso (title), odd = recto (part). Numbers centered in footer."""
    section.different_first_page_header_footer = True

    h = clear_para(section.header)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(h.add_run(recto_text), name=MONO, size=7.5, color=ASH, spacing=1.6)

    eh = clear_para(section.even_page_header)
    eh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(eh.add_run(verso_text), name=MONO, size=7.5, color=ASH, spacing=1.6)

    for footer in (section.footer, section.even_page_footer):
        fp = clear_para(footer)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if show_numbers:
            page_field(fp)

    if show_numbers:
        ffp = clear_para(section.first_page_footer)
        ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_field(ffp)
    else:
        clear_para(section.first_page_footer)
    clear_para(section.first_page_header)


# ------------------------------------------------------------- block layout

def sep_diamond(doc, before=14, after=14):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=before, after=after)
    set_font(p.add_run("◊ ◈ ◊"), name=SERIF_DISPLAY, size=11,
             color=EMBER, spacing=2)
    return p


def soft_break(doc, before=10, after=10):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=before, after=after)
    set_font(p.add_run("·  ·  ·"),
             name=SERIF_BODY, size=10.5, color=ASH)
    return p


def body_para(doc, text, first=False):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, leading=BODY_LEADING,
             indent=None if first else 0.24)
    add_inline(p, text)
    return p


def doc_block(doc, lines):
    for i, ln in enumerate(lines):
        p = para(doc, before=(8 if i == 0 else 0),
                 after=(8 if i == len(lines) - 1 else 0),
                 leading=12.6, left=0.32, right=0.22)
        if ln.strip():
            add_inline(p, ln, name=MONO, size=8.4, color=INK)
        else:
            set_font(p.add_run(" "), name=MONO, size=8.4)


def letter_block(doc, lines):
    for i, ln in enumerate(lines):
        p = para(doc, before=(8 if i == 0 else 2), after=2, leading=14.2,
                 left=0.34, right=0.3)
        add_inline(p, ln, size=10.4, base_italic=True)


def annotation_block(doc, lines):
    for i, ln in enumerate(lines):
        p = para(doc, before=(7 if i == 0 else 1), after=1, leading=11.6,
                 left=0.55, right=0.1)
        add_inline(p, ln, name=MONO, size=7.8, color=ASH)


def verse_block(doc, lines):
    for ln in lines:
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1,
                 leading=14.6)
        if ln.strip():
            add_inline(p, ln, size=10.6, base_italic=True)


def figure(doc, path, caption=None, width=4.1, plate=False):
    if plate:
        doc.add_page_break()
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=4)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    if caption:
        c = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=10)
        add_inline(c, caption, name=MONO, size=7.6, color=ASH)
    if plate:
        doc.add_page_break()


def chapter_head(doc, kicker=None, title=None, epigraph=None, epi_source=None,
                 new_page=True):
    if new_page:
        doc.add_page_break()
    para(doc, before=44, after=0)  # sink the chapter opening
    if kicker:
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, keep_next=True)
        set_font(p.add_run(kicker), name=MONO, size=8, color=EMBER,
                 spacing=2.2, caps=True)
    if title:
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, keep_next=True)
        add_inline(p, title, name=SERIF_DISPLAY, size=21, color=INK)
        p.style = doc.styles["Heading 2"] if False else p.style
    if epigraph:
        para(doc, before=4, after=0)
        for ln in epigraph:
            p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1,
                     leading=13.6, left=0.5, right=0.5)
            add_inline(p, ln, name=SERIF_DISPLAY, size=11, base_italic=True)
        if epi_source:
            p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=0)
            set_font(p.add_run(epi_source), name=MONO, size=7.6, color=ASH,
                     spacing=1.4)
    para(doc, before=0, after=16)


# --------------------------------------------------------------- the parser

def parse_chapter(doc, path):
    """Render one manuscript file into the document."""
    text = open(path, encoding="utf-8").read()
    lines = text.split("\n")
    i = 0
    kicker = title = epi_source = None
    epigraph = []
    # -- head
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("#K "):
            kicker = ln[3:].strip(); i += 1
        elif ln.startswith("#T "):
            title = ln[3:].strip(); i += 1
        elif ln.startswith("#E "):
            epigraph.append(ln[3:].strip()); i += 1
        elif ln.startswith("#ES "):
            epi_source = ln[4:].strip(); i += 1
        elif not ln.strip():
            i += 1
            if title or kicker:
                # peek: more head lines?
                j = i
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and lines[j].startswith(("#E ", "#ES ", "#K ", "#T ")):
                    i = j
                    continue
                break
        else:
            break
    chapter_head(doc, kicker=kicker, title=title,
                 epigraph=epigraph or None, epi_source=epi_source)

    # -- body
    first = True
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1
            continue
        if s.startswith("::: "):
            kind = s[4:].strip()
            i += 1
            block = []
            while i < len(lines) and lines[i].strip() != ":::":
                block.append(lines[i])
                i += 1
            i += 1  # skip closing :::
            while block and not block[0].strip():
                block.pop(0)
            while block and not block[-1].strip():
                block.pop()
            {"doc": doc_block, "letter": letter_block,
             "annotation": annotation_block, "verse": verse_block}[kind](doc, block)
            first = True
        elif s == "@sep":
            sep_diamond(doc); first = True; i += 1
        elif s == "@break":
            soft_break(doc); first = True; i += 1
        elif s == "@blank":
            para(doc, before=10, after=10); first = True; i += 1
        elif s.startswith("@fig:") or s.startswith("@plate:"):
            plate = s.startswith("@plate:")
            spec = s.split(":", 1)[1]
            fname, _, caption = spec.partition("|")
            fpath = os.path.join(BOOK, "plates", fname.strip())
            if os.path.exists(fpath):
                figure(doc, fpath, caption.strip() or None,
                       width=4.35 if plate else 4.0, plate=plate)
            first = True; i += 1
        else:
            # paragraph: consume continuous lines
            buf = [s]
            i += 1
            while i < len(lines) and lines[i].strip() and \
                    not lines[i].startswith(("::: ", "@", "#")):
                buf.append(lines[i].strip())
                i += 1
            body_para(doc, " ".join(buf), first=first)
            first = False


# ------------------------------------------------------------- book chrome

def part_opening(doc, section, number, title, plate=None, epigraph=None,
                 epi_source=None):
    para(doc, before=120, after=0)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    set_font(p.add_run(number), name=MONO, size=9.5, color=EMBER, spacing=3,
             caps=True)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_inline(p, title, name=SERIF_DISPLAY, size=30, color=INK)
    sep_diamond(doc, before=4, after=16)
    if epigraph:
        for ln in epigraph:
            p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1,
                     leading=14.4, left=0.4, right=0.4)
            add_inline(p, ln, name=SERIF_DISPLAY, size=11.5, base_italic=True)
        if epi_source:
            p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=4)
            set_font(p.add_run(epi_source), name=MONO, size=7.8, color=ASH,
                     spacing=1.4)
    if plate:
        fpath = os.path.join(BOOK, "plates", plate)
        if os.path.exists(fpath):
            doc.add_page_break()
            para(doc, before=30, after=0)
            p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
            p.add_run().add_picture(fpath, width=Inches(4.35))


def build(manifest_path, out_path):
    manifest = json.load(open(manifest_path, encoding="utf-8"))
    doc = Document()

    # global settings: even/odd headers, auto-hyphenation
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))
    if settings.find(qn("w:autoHyphenation")) is None:
        hy = OxmlElement("w:autoHyphenation")
        hy.set(qn("w:val"), "1")
        settings.append(hy)
        hz = OxmlElement("w:hyphenationZone")
        hz.set(qn("w:val"), "283")
        settings.append(hz)

    # default style guard
    normal = doc.styles["Normal"]
    normal.font.name = SERIF_BODY
    normal.font.size = Pt(BODY_PT)

    section = doc.sections[0]
    setup_section(section)
    sect_pgnum(section, fmt="lowerRoman", start=1)
    running_heads(section, manifest["running_title"], manifest["running_title"],
                  show_numbers=False)

    front = manifest.get("front_matter", [])
    for block in front:
        render_front_block(doc, block)

    for part in manifest["parts"]:
        section = doc.add_section(WD_SECTION.ODD_PAGE)
        setup_section(section)
        section.header.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        if part is manifest["parts"][0]:
            sect_pgnum(section, fmt="decimal", start=1)
        else:
            sect_pgnum(section, fmt="decimal")
        running_heads(section, manifest["running_title"],
                      part.get("running", part["title"]).upper(), True)
        part_opening(doc, section, part["number"], part["title"],
                     plate=part.get("plate"), epigraph=part.get("epigraph"),
                     epi_source=part.get("epi_source"))
        for ch in part["chapters"]:
            parse_chapter(doc, os.path.join(BOOK, "manuscript", ch))

    for block in manifest.get("back_matter", []):
        section = doc.add_section(WD_SECTION.ODD_PAGE)
        setup_section(section)
        for hf in (section.header, section.even_page_header, section.footer,
                   section.even_page_footer, section.first_page_header,
                   section.first_page_footer):
            hf.is_linked_to_previous = False
        sect_pgnum(section, fmt="decimal")
        running_heads(section, manifest["running_title"],
                      block.get("running", "").upper() or
                      manifest["running_title"], True)
        parse_chapter(doc, os.path.join(BOOK, "manuscript", block["file"]))

    doc.save(out_path)
    print(f"saved {out_path}")


# ---------------------------------------------------- front matter renderers

def render_front_block(doc, block):
    t = block["type"]
    if t == "page_break":
        doc.add_page_break()
    elif t == "file":
        parse_front_file(doc, os.path.join(BOOK, "manuscript", block["file"]),
                         new_page=block.get("new_page", True))
    elif t == "plate":
        doc.add_page_break()
        para(doc, before=40, after=0)
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
        fpath = os.path.join(BOOK, "plates", block["file"])
        if os.path.exists(fpath):
            p.add_run().add_picture(fpath, width=Inches(block.get("width", 4.35)))
        if block.get("caption"):
            c = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=4)
            add_inline(c, block["caption"], name=MONO, size=7.6, color=ASH)


def parse_front_file(doc, path, new_page=True):
    """Front-matter files reuse the chapter markup but without page-break-and-
    sink chapter heads; they manage their own composition via directives."""
    text = open(path, encoding="utf-8").read()
    if new_page:
        doc.add_page_break()
    lines = text.split("\n")
    i = 0
    first = True
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1
            continue
        if s.startswith("!space:"):
            para(doc, before=int(s.split(":")[1]), after=0); i += 1
        elif s.startswith("!center:"):
            payload = ln.split(":", 1)[1]
            style, _, txt = payload.partition("|")
            p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2,
                     leading=15)
            render_styled(p, style, txt)
            i += 1
        elif s.startswith("!left:"):
            payload = ln.split(":", 1)[1]
            style, _, txt = payload.partition("|")
            p = para(doc, before=2, after=2, leading=14)
            render_styled(p, style, txt)
            i += 1
        elif s == "@sep":
            sep_diamond(doc); i += 1
        elif s == "@break":
            soft_break(doc); i += 1
        elif s.startswith("::: "):
            kind = s[4:].strip()
            i += 1
            block = []
            while i < len(lines) and lines[i].strip() != ":::":
                block.append(lines[i]); i += 1
            i += 1
            {"doc": doc_block, "letter": letter_block,
             "annotation": annotation_block, "verse": verse_block}[kind](doc, block)
        elif s == "@pagebreak":
            doc.add_page_break(); first = True; i += 1
        else:
            buf = [s]; i += 1
            while i < len(lines) and lines[i].strip() and \
                    not lines[i].startswith(("!", "@", ":::")):
                buf.append(lines[i].strip()); i += 1
            body_para(doc, " ".join(buf), first=first)
            first = False


def render_styled(p, style, txt):
    styles = {
        "title":    dict(name=SERIF_DISPLAY, size=34, color=INK),
        "subtitle": dict(name=SERIF_DISPLAY, size=15, color=INK),
        "display":  dict(name=SERIF_DISPLAY, size=13, color=INK),
        "kicker":   dict(name=MONO, size=8.5, color=EMBER, spacing=2.4, caps=True),
        "mono":     dict(name=MONO, size=8.4, color=INK),
        "monosm":   dict(name=MONO, size=7.6, color=ASH),
        "body":     dict(name=SERIF_BODY, size=BODY_PT, color=INK),
        "bodysm":   dict(name=SERIF_BODY, size=9, color=ASH),
        "italic":   dict(name=SERIF_BODY, size=BODY_PT, color=INK),
        "ember":    dict(name=SERIF_DISPLAY, size=12, color=EMBER),
    }
    kw = styles.get(style, styles["body"])
    base_italic = style == "italic"
    add_inline(p, txt, base_italic=base_italic, **kw)


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else os.path.join(HERE, "manifest.json"),
          sys.argv[2] if len(sys.argv) > 2 else os.path.join(BOOK, "USING_THIS_MATTER.docx"))
